"""
Extracts plain text from uploaded documents so it can be chunked and embedded.

Kept dependency-light and synchronous-under-the-hood (PDF/DOCX parsing
libraries are CPU-bound, not I/O-bound) but exposed via an async wrapper that
runs the parse in a thread so it never blocks the event loop.
"""
import io
from pathlib import Path

import anyio

from app.core.exceptions import UnsupportedFileTypeError, ValidationFailedError
from app.core.logging_config import get_logger

logger = get_logger(__name__)


def _extract_pdf(content: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(text)
    return "\n\n".join(pages)


def _extract_docx(content: bytes) -> str:
    import docx  # python-docx

    document = docx.Document(io.BytesIO(content))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)
    return "\n".join(parts)


def _extract_plain_text(content: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValidationFailedError("Could not decode text file with any supported encoding.")


_EXTRACTORS = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".txt": _extract_plain_text,
    ".md": _extract_plain_text,
}


def _extract_sync(filename: str, content: bytes) -> str:
    ext = Path(filename).suffix.lower()
    extractor = _EXTRACTORS.get(ext)
    if extractor is None:
        raise UnsupportedFileTypeError(f"Unsupported file type '{ext}'. Allowed: {', '.join(_EXTRACTORS)}")

    try:
        text = extractor(content)
    except UnsupportedFileTypeError:
        raise
    except Exception as exc:
        logger.error("Text extraction failed", extra={"context": {"filename": filename, "error": str(exc)}})
        raise ValidationFailedError(f"Could not extract text from '{filename}': {exc}") from exc

    text = text.strip()
    if not text:
        raise ValidationFailedError(
            f"No extractable text found in '{filename}'. It may be a scanned/image-only document."
        )
    return text


async def extract_text(filename: str, content: bytes) -> str:
    """Runs the (blocking) extraction logic in a worker thread."""
    return await anyio.to_thread.run_sync(_extract_sync, filename, content)
